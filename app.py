from flask import Flask, request, send_file
from docx import Document
import os
import io
import tempfile
from flask_cors import CORS
from datetime import datetime

app = Flask(__name__)
CORS(app)

def format_date(date_str):
    try:
        return datetime.strptime(date_str, "%d.%m.%Y").strftime("%d.%m.%Y")
    except:
        return date_str

def replace_placeholders_in_paragraph(paragraph, placeholders):
    full_text = ''.join(run.text for run in paragraph.runs)
    for key, val in placeholders.items():
        if key in full_text:
            full_text = full_text.replace(key, val)
    for run in paragraph.runs:
        run.text = ''
    if paragraph.runs:
        paragraph.runs[0].text = full_text

def replace_in_doc(doc, placeholders):
    for para in doc.paragraphs:
        replace_placeholders_in_paragraph(para, placeholders)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    replace_placeholders_in_paragraph(para, placeholders)

@app.route('/')
def index():
    return "DOCX contract generator for gas – FIRMA is running."

@app.route('/api/generate', methods=['POST'])
def generate():
    data = request.json

    doc = Document("Rekapitulace_plyn_firma.docx")
    placeholders = {
        "{{cislo_smlouvy}}": data.get("cislo_smlouvy", ""),
        "{{cislo_partnera}}": data.get("cislo_partnera", ""),
        "{{Nazev}}": data.get("firma", ""),
        "{{ICO}}": data.get("ico", ""),
        "{{ulice_sidlo}}": data.get("ulice_sidlo", ""),
        "{{mesto_sidlo}}": data.get("mesto_sidlo", ""),
        "{{psc_sidlo}}": data.get("psc_sidlo", ""),
        "{{email}}": data.get("email", ""),
        "{{telefon}}": data.get("telefon", ""),
        "{{zpusob_odesilani}}": data.get("zpusob_odesilani", ""),
        "{{platby_faktury}}": data.get("platby_faktury", ""),
        "{{platby_zalohy}}": data.get("platby_zalohy", ""),
        "{{cislo_uctu}}": data.get("cislo_uctu", ""),
        "{{zahajeni_dodavek}}": format_date(data.get("zahajeni_dodavek", "")),
        "{{prolongace}}": format_date(data.get("prolongace", "")),
        "{{ean}}": data.get("ean", ""),
        "{{ulice_odber}}": data.get("ulice_odber", ""),
        "{{mesto_odber}}": data.get("mesto_odber", ""),
        "{{psc_odber}}": data.get("psc_odber", "")
    }

    replace_in_doc(doc, placeholders)

    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
        doc.save(tmp.name)
        tmp.seek(0)
        return send_file(tmp.name, as_attachment=True, download_name="Rekapitulace_smlouvy_firma_plyn.docx")

if __name__ == "__main__":
    app.run(debug=False, host="0.0.0.0", port=10000)
